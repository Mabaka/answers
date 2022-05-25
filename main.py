import re
import requests
from bs4 import BeautifulSoup
import numpy as numpy
import pandas
from docx import Document

def handle_query(result):
    links = []
    titles = []
    result_array = []
    for g in result.find_all('div', class_='yuRUbf'):
        anchors = g.find_all('a')
        if anchors:
            link = anchors[0]['href']
            title = g.find('h3').text

            links.append(link)
            titles.append(title)

    result_array.append(links)
    result_array.append(titles)

    result_array_np = numpy.array(result_array)
    result_array_pd = pandas.DataFrame(result_array_np.transpose(1,0), columns=['link','title'])

    return result_array_pd

def handle_links(links,titles):
    answer = ''
    i = 0
    for link in links:
        URL = link
        headers = {"user-agent": USER_AGENT}
        try:
            resp = requests.get(URL, headers=headers)
            if resp.status_code == 200:
                soup = BeautifulSoup(resp.content, "html.parser")
                all_p = soup.find_all('p')
                for p in all_p:
                    if len(answer) > 2000:
                        break
                    answer+= handle_p(titles[i],p)
                i+=1
        except Exception as e:
            print(e)
            i += 1

    return answer

def handle_p(title,p):
    text_answer = ''
    title = re.sub(r"\d+","",title,flags=re.UNICODE)
    title = re.sub(r'[^\w\s]', '', title)
    if len(p.text) > 10:

        key_words = title.split(' ')
        for key_word in key_words:
            if p.text.find(key_word) != -1 and len(key_word) > 4 \
                    and p.text.lower().find('содержание') == -1:
                text_answer_ar = p.text.split('.')[:5]
                if len(text_answer_ar) > 0:
                    for c in range(len(text_answer_ar)):
                        if len(text_answer_ar[c]) > 300:
                            continue
                        text_answer+= text_answer_ar[c]
                else:
                    text_answer += p.text
                break
    text_answer = re.sub('<[^>]*>', '', text_answer)
    text_answer = re.sub(r"\b[^а-яА-Я]\n^\s\b", "", text_answer)
    text_answer = text_answer.replace(f'\n',' ')
    return text_answer

USER_AGENT = "Mozilla/5.0 (Macintosh; Intel Mac OS X 10.14; rv:65.0) Gecko/20100101 Firefox/65.0"
QUESTIONS_FILE = open('questions.txt','r')
QUESTIONS_TEXT = QUESTIONS_FILE.read()
QUESTIONS_ARRAY = QUESTIONS_TEXT.split('#')
QUESTIONS_ARRAY_NUMPY = numpy.array(QUESTIONS_ARRAY)
QUESTIONS_DF = pandas.DataFrame(QUESTIONS_ARRAY,columns=['question'])

questions_df_copy = QUESTIONS_DF.copy()
answers = []

for question in QUESTIONS_DF['question']:

    query = question
    query = query.replace(' ','+')
    URL = f'https://google.com/search?q={query}'

    headers = {"user-agent": USER_AGENT}
    resp = requests.get(URL, headers=headers)

    if resp.status_code == 200:
        soup = BeautifulSoup(resp.content, "html.parser")
        links_df = handle_query(soup)
        answer = handle_links(links_df['link'],links_df['title'])
        answers.append(answer)

answers_np = numpy.array(answers)
questions_df_copy['answer'] = answers_np

doc = Document()
q_len = len(questions_df_copy)

for i in range(q_len):
    doc.add_paragraph(f'{i+1} ' + questions_df_copy.iloc[i]['question'])
    doc.add_paragraph(questions_df_copy.iloc[i]['answer'])

doc.save('Otvety.docx')



